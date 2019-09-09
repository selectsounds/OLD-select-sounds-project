import os
import csv
from typing import List

import docx
from docx.opc.exceptions import PackageNotFoundError

from src import config

DIR_PATH = os.path.abspath(os.path.dirname(__file__))


def setup_doc_settings() -> docx.Document:
    try:
        labels_doc = docx.Document(config.ROOT_DIR + '/data/labels_template.docx')
    except PackageNotFoundError as err:
        return err
    style = labels_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(8)
    font.bold = True

    return labels_doc


def get_label_cells(document: docx.Document) -> List:
    tables = [t for t in document.tables]

    rows = [r for t in tables for r in t.rows]

    cells = [c for r in rows for c in r.cells]

    return cells


def read_records_csv_file() -> List:
    records = []
    with open(config.ROOT_DIR + '/data/records.csv') as record_csv_file:
        csv_data = csv.reader(record_csv_file, delimiter=',')
        # print(csv_data)

        records = [r for r in csv_data][1:]

    return records


def format_record_data(record_data: List) -> List:
    record_text = [
        "\n".join([
            f'Album: {r[0]}', f'Artist: {r[1]}', f'Label: {r[2]}', f'Country: {r[3]}',
            f'Release date: {r[4]}', f'Speed: {r[5]}',
            'Tracklist: {}'.format(r[6][2:-2].replace('"', '').replace(',', ', ')),
        ])
        for r in record_data
    ]

    return record_text


def save_labels(labels_doc: docx.Document, path=None) -> None:
    if not path:
        path = config.ROOT_DIR + '/data/labels.docx'

    labels_doc.save(path)


if __name__ == '__main__':
    test_doc = docx.Document(config.ROOT_DIR + '/data/labels_template.docx')


def check_labels_file_exists():
    file_exists = os.path.exists(config.ROOT_DIR + '/data/labels.docx')
    return file_exists
